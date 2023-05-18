import glob,random,re,json,sys
from docx import Document


max_seq_len= 512

def get_docx_textlines(doc, table_ignore_flag=True):
    raw_text = ''
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        #text = text.replace(' ','')
        text = ''.join(text.split())
        if len(text) > 0 :
            raw_text += text + '\n'
    #print (raw_text)

    if not table_ignore_flag:
        tables = doc.tables
        for tables_i in tables:
            text = ''
            for row in tables_i.rows:
                row_content = []
                for cell in row.cells:
                    if cell.text not in row_content:
                        row_content.append(cell.text)
                row_text = '\t\t'.join(row_content)
                text += row_text + '\n'
            raw_text += text

    return raw_text

def encode_prompt(instructions, inputs):
    input = random.sample(inputs,1)[0]
    instruction = random.sample(instructions,1)[0]
    prompt = {}
    prompt['instruction'] = instruction
    prompt['input'] = input
    return prompt

instructions_raw = [
    '请对输入文本进行段落划分。输出每个段落的序号以及每个段落包含的原始文本句子，示例：段落1： 段落2： 。'
]

with open('prompts_contract_raw.json','w',encoding='utf-8') as res:
    docs = glob.glob('./contract/*.docx')
    prompts = []
    for doc in docs:
        word_raw = Document(doc)
        raw_text = get_docx_textlines(word_raw)
        sentences = raw_text.split('\n')
        # 合并
        merged_sentences = []
        raw_sentences = []
        start_index_ = 0
        while start_index_ < len(sentences):
            tmp_text = sentences[start_index_]
            end_index_ = start_index_ + 1
            while end_index_ < len(sentences) and \
                    len(tmp_text) + len(sentences[end_index_]) <= max_seq_len - 2:
                tmp_text += '\n' + sentences[end_index_]
                end_index_ += 1
            start_index_ = end_index_
            merged_sentences.append(tmp_text)
            raw_sentences.append(tmp_text.replace('\n',''))
        for idx,input in enumerate(merged_sentences):
            prompt = encode_prompt(instructions_raw, [input])
            prompts.append(prompt)
    print (len(prompts))
    json.dump(prompts,res,ensure_ascii=False)
res.close()


